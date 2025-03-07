import streamlit as st
import requests
import groq
import json
import urllib.parse
import feedparser
from googleapiclient.discovery import build
from docx import Document
from io import BytesIO
from docx.shared import Pt

# Blacklisted domains
BLACKLISTED_DOMAINS = ["coursera.org", "udemy.com", "edx.org", "skillshare.com", "khanacademy.org", "simplilearn.com", "geeksforgeeks"]

def filter_links(links):
    """Filters out links that belong to blacklisted domains"""
    return [link for link in links if not any(domain in link for domain in BLACKLISTED_DOMAINS)]



# Streamlit UI
st.title("🔍 Advanced Research Assistant")

# # User Input for API Keys
# if "GROQ_API_KEY" not in st.session_state:
#     st.session_state.GROQ_API_KEY = st.text_input("Enter your Groq API Key:", type="password")

# if "YOUTUBE_API_KEY" not in st.session_state:
#     st.session_state.YOUTUBE_API_KEY = st.text_input("Enter your YouTube API Key:", type="password")

# if "SERPAPI_KEY" not in st.session_state:
#     st.session_state.SERPAPI_KEY = st.text_input("Enter your SerpAPI Key:", type="password")

# # Ensure all API keys are entered before proceeding
# if not all([st.session_state.GROQ_API_KEY, st.session_state.YOUTUBE_API_KEY, st.session_state.SERPAPI_KEY]):
#     st.warning("Please enter all API keys to continue.")
#     st.stop()

# Initialize session state variables if they are not already set
if "GROQ_API_KEY" not in st.session_state:
    st.session_state.GROQ_API_KEY = ""

if "YOUTUBE_API_KEY" not in st.session_state:
    st.session_state.YOUTUBE_API_KEY = ""

if "SERPAPI_KEY" not in st.session_state:
    st.session_state.SERPAPI_KEY = ""

# Function to update session state without causing a full reload
def update_api_key(key, value):
    st.session_state[key] = value

# User input fields with default values from session state
st.session_state.GROQ_API_KEY = st.text_input("Enter your Groq API Key:", 
                                              value=st.session_state.GROQ_API_KEY, 
                                              type="password", 
                                              on_change=update_api_key, 
                                              args=("GROQ_API_KEY", st.session_state.GROQ_API_KEY))

st.session_state.YOUTUBE_API_KEY = st.text_input("Enter your YouTube API Key:", 
                                                 value=st.session_state.YOUTUBE_API_KEY, 
                                                 type="password", 
                                                 on_change=update_api_key, 
                                                 args=("YOUTUBE_API_KEY", st.session_state.YOUTUBE_API_KEY))

st.session_state.SERPAPI_KEY = st.text_input("Enter your SerpAPI Key:", 
                                             value=st.session_state.SERPAPI_KEY, 
                                             type="password", 
                                             on_change=update_api_key, 
                                             args=("SERPAPI_KEY", st.session_state.SERPAPI_KEY))

# Initialize Groq Client
groq_client = groq.Groq(api_key=st.session_state.GROQ_API_KEY)

def get_ai_summary(topic):
    """Generates an AI-based summary using Groq API"""
    response = groq_client.chat.completions.create(
        model="mixtral-8x7b-32768",
        messages=[
            {"role": "system", "content": "You are a research assistant."},
            {"role": "user", "content": f"Provide a summary of {topic}."}
        ]
    )
    return response.choices[0].message.content

def fetch_research_papers_google_scholar(query, num_pages=2):
    """Fetches research papers using Google Scholar via SerpAPI"""
    papers = []

    for page in range(num_pages):
        params = {
            "engine": "google_scholar",
            "q": query,
            "start": page * 10,  # Pagination (10 results per page)
            "api_key": st.session_state.SERPAPI_KEY
        }
        response = requests.get("https://serpapi.com/search", params=params)

        if response.status_code == 200:
            results = response.json().get("organic_results", [])
            for result in results:
                title = result.get("title", "No Title")
                link = result.get("link", "#")
                snippet = result.get("snippet", "No summary available")

                if not any(domain in link for domain in BLACKLISTED_DOMAINS):
                    papers.append(f"Title: {title}\nSummary: {snippet[:300]}...\n[Read More]({link})\n\n")

    return papers if papers else ["No relevant research papers found."]

def get_youtube_videos(query):
    """Fetches YouTube videos using Google API"""
    youtube = build("youtube", "v3", developerKey=st.session_state.YOUTUBE_API_KEY)

    request = youtube.search().list(
        q=query,
        part="id,snippet",
        maxResults=5,
        order="relevance"  # Sort by most viewed videos
    )

    response = request.execute()
    videos = []

    for item in response.get("items", []):
        if "videoId" in item.get("id", {}):
            video_title = item["snippet"]["title"]
            video_url = f"https://www.youtube.com/watch?v={item['id']['videoId']}"

            if not any(edtech in video_title.lower() for edtech in BLACKLISTED_DOMAINS):
                videos.append(f"[{video_title}]({video_url})")
    return videos if videos else ["No relevant videos found."]

def get_resources(topic):
    ai_summary = get_ai_summary(topic)
    research_papers = fetch_research_papers_google_scholar(topic)
    youtube_videos = get_youtube_videos(topic)
    return ai_summary, research_papers, youtube_videos

# Without Proper same name of document
# def create_word_doc(topic, ai_summary, research_papers, youtube_videos):
#     """Creates a downloadable Word document with research content"""
#     doc = Document()
#     doc.add_heading(f"Research Report on {topic}", level=1)

#     doc.add_heading("AI Summary", level=2)
#     doc.add_paragraph(ai_summary)

#     doc.add_heading("Latest Research Papers", level=2)
#     for paper in research_papers:
#         doc.add_paragraph(paper)

#     doc.add_heading("YouTube Videos", level=2)
#     for video in youtube_videos:
#         doc.add_paragraph(video)

#     # Save document to in-memory file
#     file_stream = BytesIO()
#     doc.save(file_stream)
#     file_stream.seek(0)
#     return file_stream

# Code for Dynamic Heading of topics
# def create_word_doc(topic, ai_summary, research_papers, youtube_videos):
#     """Creates a downloadable Word document with research content"""
#     doc = Document()
#     doc.add_heading(f"{topic} - Summary", level=1)  # Change heading here

#     doc.add_paragraph(ai_summary)

#     doc.add_heading("Latest Research Papers", level=2)
#     for paper in research_papers:
#         doc.add_paragraph(paper)

#     doc.add_heading("YouTube Videos", level=2)
#     for video in youtube_videos:
#         doc.add_paragraph(video)

#     # Save document to in-memory file
#     file_stream = BytesIO()
#     doc.save(file_stream)
#     file_stream.seek(0)
#     return file_stream

# def create_word_doc(topic, ai_summary, research_papers, youtube_videos):
#     """Creates a downloadable Word document with research content"""
#     doc = Document()

#     # Title (Topic + Summary) - Bold and Large
#     title = doc.add_heading(level=1)
#     run = title.add_run(f"{topic} - Summary")
#     run.bold = True
#     run.font.size = Pt(16)  # Increase size

#     # Summary Section
#     doc.add_paragraph(ai_summary)

#     # Latest Research Papers (Bold & Larger)
#     research_heading = doc.add_heading(level=2)
#     run = research_heading.add_run("Latest Research Papers")
#     run.bold = True
#     run.font.size = Pt(14)

#     for paper in research_papers:
#         doc.add_paragraph(paper)
    # # YouTube Videos Section (Bold & Larger)
    # youtube_heading = doc.add_heading(level=2)
    # run = youtube_heading.add_run("YouTube Videos")
    # run.bold = True
    # run.font.size = Pt(14)

    # for video in youtube_videos:
    #     doc.add_paragraph(video)

    # # Save document to in-memory file
    # file_stream = BytesIO()
    # doc.save(file_stream)
    # file_stream.seek(0)
    # return file_stream

def create_word_doc(topic, ai_summary, research_papers, youtube_videos):
    """Creates a downloadable Word document with research content"""
    doc = Document()

    # Add title
    title = doc.add_paragraph()
    title_run = title.add_run(f"Research Report on {topic}")
    title_run.bold = True
    title_run.font.size = Pt(20)

    # Add a blank line for spacing
    doc.add_paragraph()

    # Add Summary Heading
    summary_heading = doc.add_paragraph()
    summary_run = summary_heading.add_run(f"{topic} - Summary")
    summary_run.bold = True
    summary_run.font.size = Pt(16)

    # Add Summary Content
    #doc.add_paragraph(ai_summary)
    summary_paragraph = doc.add_paragraph()
    summary_paragraph.add_run(ai_summary)

    # Add a blank line for spacing
    doc.add_paragraph()

    # Add Research Papers Heading
    research_heading = doc.add_paragraph()
    research_run = research_heading.add_run("Latest Research Papers")
    research_run.bold = True
    research_run.font.size = Pt(16)

    for paper in research_papers:
        paper_paragraph = doc.add_paragraph()
        paper_paragraph.add_run(paper)

    # Add a blank line for spacing
    doc.add_paragraph()

    # Add YouTube Videos Heading
    video_heading = doc.add_paragraph()
    video_run = video_heading.add_run("YouTube Videos")
    video_run.bold = True
    video_run.font.size = Pt(16)


    for video in youtube_videos:
        video_paragraph = doc.add_paragraph()
        video_paragraph.add_run(video)

    # Save document to in-memory file
    from io import BytesIO
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


# User input
topic = st.text_input("Enter a research topic:")

if st.button("Get Research Resources"):
    if topic:
        ai_summary, research_papers, youtube_videos = get_resources(topic)
        
        st.markdown(f"## 📌 **{topic} - Summary**")  # Bold title
        st.write(ai_summary)

        st.markdown("## 📖 **Latest Research Papers**")  # Bold section title
        for paper in research_papers:
            st.markdown(paper)

        st.markdown("## 🎥 **YouTube Videos**")  # Bold section title
        for video in youtube_videos:
            st.markdown(video)


        # Create downloadable Word document
        word_file = create_word_doc(topic, ai_summary, research_papers, youtube_videos)
        st.download_button(label="📥 Download Report (Word)",
                           data=word_file,
                           file_name=f"{topic}_Research_Report.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.warning("Please enter a topic before searching.")
